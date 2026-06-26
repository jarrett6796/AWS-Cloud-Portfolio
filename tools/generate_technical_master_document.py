from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_OUT = ROOT / "TECHNICAL_MASTER_DOCUMENT.md"
DOCX_OUT = ROOT / "Multi-Cloud AI Portfolio Assistant 技術文件.docx"


TITLE = "Multi-Cloud AI Portfolio Assistant 工程技術文件"
SUBTITLE = "Google Docs 版技術說明文件｜根據目前 repo Markdown、frontend-AWS、backend-GCP 與 GitHub Actions 內容整理"


SECTIONS = [
    {
        "heading": "1. 背景與動機",
        "body": [
            (
                "本專案從 Cloud Resume Challenge 與 AWS Cloud Engineer 轉職學習路線出發。"
                "傳統履歷可以列出技能，但很難完整呈現雲端資源、前後端整合、部署流程、"
                "除錯紀錄與架構取捨。因此，專案一開始以 AWS 雲端履歷網站與 serverless "
                "visitor counter 作為可驗證的實作起點。"
            ),
            (
                "在後續迭代中，專案從靜態作品集演進成互動式 AI Portfolio。訪客不只可以閱讀"
                "頁面上的專案說明，也可以透過 AI Assistant 用自然語言詢問架構、部署、RAG、"
                "troubleshooting 與技術決策。"
            ),
        ],
        "subsections": [
            {
                "heading": "1.1 背景故事",
                "body": [
                    (
                        "我正在學習雲端技術並轉職 AWS Cloud Engineer。專案起點是 Cloud Resume "
                        "Challenge：先建立可部署、可公開訪問、可被驗證的雲端履歷網站，再逐步加入"
                        "真實工程功能。"
                    )
                ],
            },
            {
                "heading": "1.2 問題發現",
                "bullets": [
                    "靜態履歷缺乏互動性，無法直接展示雲端系統如何運作。",
                    "GitHub、技術文件、部署紀錄與作品集內容分散，面試官需要自行拼出完整脈絡。",
                    "傳統作品集通常只展示結果畫面，較難呈現 root cause 分析、CORS、Cloud Run revision、CI/CD env var 等工程細節。",
                    "求職者很難只靠履歷條列展現 cloud engineering、backend、frontend、AI/RAG 與 troubleshooting 的整合能力。",
                ],
            },
            {
                "heading": "1.3 解決方案",
                "bullets": [
                    "建立 React + Vite portfolio website，並透過 AWS S3 + CloudFront 對外提供 HTTPS/CDN 存取。",
                    "加入 AWS serverless visitor counter：frontend 呼叫 API Gateway endpoint，由 Lambda 與 DynamoDB 處理瀏覽次數。",
                    "導入 GCP Cloud Run FastAPI backend 作為 AI/RAG backend，使用 Gemini、Firestore 與 Cloud Storage 建立可查詢的知識庫。",
                    "AI Assistant 以 `/ask-rag-stream` 作為主要路徑，支援 streaming response、來源 metadata 與 `/ask-rag` fallback。",
                ],
            },
            {
                "heading": "1.4 啟發與動機",
                "body": [
                    (
                        "這個專案的重點不是把所有雲端服務堆在一起，而是把一個原本單一 AWS "
                        "Cloud Resume 專案，演進成可以展示 Multi-Cloud、AI Engineering、Backend "
                        "API、Frontend UX、部署流程與實際除錯能力的工程作品。"
                    )
                ],
            },
        ],
    },
    {
        "heading": "2. 系統概覽",
        "subsections": [
            {
                "heading": "2.1 專案介紹",
                "body": [
                    (
                        "Multi-Cloud AI Portfolio Assistant 是一個雲端工程作品集平台。前端 portfolio "
                        "部署於 AWS 靜態網站/CDN 路徑，visitor counter 使用 AWS serverless 架構；"
                        "AI assistant 則使用 GCP Cloud Run 上的 FastAPI backend，透過 Firestore 中的 "
                        "document chunks、conversation history 與 Vertex AI Gemini 產生 grounded answers。"
                    )
                ],
            },
            {
                "heading": "2.2 核心功能",
                "table": {
                    "headers": ["功能", "說明", "使用技術"],
                    "rows": [
                        ["Portfolio Website", "公開展示個人背景、技能、作品集與 capstone project。", "React, Vite, JavaScript, CSS, S3, CloudFront"],
                        ["Visitor Counter", "顯示 live view count，frontend 呼叫已部署的 `/views` endpoint。", "API Gateway, Lambda, DynamoDB"],
                        ["AI Assistant", "讓訪客以自然語言詢問專案內容與架構。", "React ChatPanel, Cloud Run, FastAPI, Gemini"],
                        ["RAG Knowledge Base", "將專案 markdown 文件切 chunk、產生 embedding、存入 Firestore 後檢索。", "GCS, Firestore, text-embedding-005, cosine similarity"],
                        ["Conversation History", "以 `session_id` 保存 user/assistant 對話，用於 follow-up context。", "Firestore `conversations/{session_id}/messages/{message_id}`"],
                        ["Streaming Response", "AI 回答以 SSE metadata/token/done/error 事件逐步回傳。", "POST `/ask-rag-stream`, ReadableStream, Server-Sent Events"],
                    ],
                },
            },
            {
                "heading": "2.3 技術棧",
                "table": {
                    "headers": ["分類", "技術", "用途"],
                    "rows": [
                        ["Frontend", "React 19, Vite 8, JavaScript, CSS", "Portfolio UI、bilingual content、theme、modal、AI assistant。"],
                        ["AWS", "S3, CloudFront, API Gateway, Lambda, DynamoDB", "靜態網站託管/CDN/HTTPS 與 visitor counter serverless path。"],
                        ["GCP", "Cloud Run, Cloud Storage, Firestore, Artifact Registry", "RAG backend hosting、文件來源、chunk/conversation 儲存與 container deployment。"],
                        ["Backend", "Python 3.11, FastAPI, Uvicorn", "API routes、RAG orchestration、health checks、streaming response。"],
                        ["AI/RAG", "Gemini 2.5 Flash, text-embedding-005", "回答生成、embedding、retrieval context 建立。"],
                        ["Database", "Firestore, DynamoDB", "RAG chunks/conversations 與 visitor count。"],
                        ["DevOps", "GitHub Actions, Docker, gcloud, AWS CLI", "frontend S3 sync/CloudFront invalidation 與 backend Cloud Run deployment。"],
                    ],
                },
            },
        ],
    },
    {
        "heading": "3. 架構設計",
        "subsections": [
            {
                "heading": "3.1 使用者服務藍圖",
                "numbered": [
                    "使用者進入 CloudFront 提供的 Portfolio Website。",
                    "使用者查看 portfolio、capstone card、project modal 與 documentation hub。",
                    "Frontend 呼叫 AWS API Gateway `/views` endpoint 取得 visitor count。",
                    "使用者開啟 AI Assistant 並輸入專案問題。",
                    "Frontend 先呼叫 GCP Cloud Run backend 的 `/ask-rag-stream`。",
                    "Backend 讀取 Firestore conversation history，視設定執行 query rewriting，產生 embedding 並檢索 Firestore `document_chunks`。",
                    "Gemini 根據 retrieved context 產生回答，frontend 逐步顯示 streaming tokens 與 sources。",
                ],
            },
            {
                "heading": "3.2 AWS 與 GCP 整體架構圖",
                "code": """```mermaid
flowchart LR
  User[User / Recruiter / Interviewer] --> CF[AWS CloudFront]
  CF --> S3[S3 Static Frontend]
  S3 --> React[React + Vite Portfolio]
  React --> APIGW[API Gateway /views]
  APIGW --> Lambda[Lambda Visitor Counter]
  Lambda --> DDB[DynamoDB View Count]
  React --> CR[Cloud Run FastAPI Backend]
  CR --> FSChunks[Firestore document_chunks]
  CR --> FSConvos[Firestore conversations]
  CR --> GCS[Cloud Storage Markdown Sources]
  CR --> Vertex[Vertex AI Gemini + text-embedding-005]
  Vertex --> CR
  CR --> React
```""",
            },
            {
                "heading": "3.2.1 AWS 架構圖",
                "body": [
                    (
                        "AWS 負責 portfolio delivery 與 visitor counter。Frontend build output 位於 `frontend-AWS/dist/`，"
                        "GitHub Actions 透過 `aws s3 sync dist/ s3://... --delete` 部署到 S3，並建立 CloudFront invalidation。"
                    )
                ],
                "code": """```mermaid
flowchart LR
  User --> CloudFront
  CloudFront --> S3[S3 static website assets]
  S3 --> Browser[React app in browser]
  Browser --> APIGateway[API Gateway /views]
  APIGateway --> Lambda
  Lambda --> DynamoDB
  DynamoDB --> Lambda --> APIGateway --> Browser
```""",
            },
            {
                "heading": "3.2.2 GCP 架構圖",
                "body": [
                    (
                        "GCP 負責 AI/RAG backend。Cloud Run 執行 FastAPI service，讀取 Cloud Storage 中的 markdown 文件，"
                        "將 chunks 與 embeddings 寫入 Firestore，並用 Gemini 2.5 Flash 產生回答。"
                    )
                ],
                "code": """```mermaid
flowchart LR
  React[React ChatPanel] --> Stream[POST /ask-rag-stream]
  Stream --> CloudRun[Cloud Run FastAPI]
  CloudRun --> History[Firestore conversations]
  CloudRun --> Chunks[Firestore document_chunks]
  CloudRun --> GCS[Cloud Storage docs]
  CloudRun --> Embedding[text-embedding-005]
  CloudRun --> Gemini[Gemini 2.5 Flash]
  Gemini --> CloudRun --> React
```""",
            },
            {
                "heading": "3.3 系統模組圖",
                "table": {
                    "headers": ["模組", "責任", "主要檔案或服務"],
                    "rows": [
                        ["Frontend UI", "Portfolio、navbar、modal、AI assistant、language/theme state。", "`frontend-AWS/src/pages/Home.jsx`, `components/`, `hooks/`"],
                        ["AWS Visitor Counter", "取得與顯示瀏覽次數。", "`frontend-AWS/src/api/visitors.js`, API Gateway, Lambda, DynamoDB"],
                        ["GCP AI Backend", "API routing、CORS、health、RAG orchestration、streaming。", "`backend-GCP/main.py`, `app/routes/`, Cloud Run"],
                        ["RAG Retrieval Layer", "chunking、embedding、cosine similarity、hybrid scoring、reranking。", "`app/services/vector_service.py`, `rag_service.py`"],
                        ["Conversation Storage", "保存 session-based user/assistant messages 與 query-rewrite audit。", "Firestore `conversations/{session_id}/messages/{message_id}`"],
                        ["Document Storage", "保存 source markdown documents 並提供 ingestion 來源。", "Cloud Storage bucket `cloud-resume-ai-rag-docs`"],
                    ],
                },
            },
            {
                "heading": "3.4 資料流程圖",
                "code": """```mermaid
sequenceDiagram
  participant U as User
  participant F as React Frontend
  participant B as Cloud Run FastAPI
  participant H as Firestore conversations
  participant C as Firestore document_chunks
  participant V as Vertex AI / Gemini

  U->>F: Ask a project question
  F->>B: POST /ask-rag-stream with question + session_id
  B->>H: Load recent visible history
  B->>V: Optional query rewrite and embedding
  B->>C: Stream chunks and score retrieval candidates
  B->>V: Generate grounded answer from retrieved context
  B-->>F: SSE metadata with sources
  B-->>F: SSE token events
  B->>H: Save user, optional system audit, assistant message
  B-->>F: SSE done
  F-->>U: Render answer + sources
```""",
            },
        ],
    },
    {
        "heading": "4. 實作流程",
        "subsections": [
            {
                "heading": "4.1 AWS 前後端",
                "body": [
                    "前端實作位於 `frontend-AWS/`，使用 `npm run build` 產出 `dist/`。部署 workflow 會安裝 dependencies、建立 `.env`、build frontend、同步到 S3，最後 invalidates CloudFront cache。",
                    "Visitor counter 的 frontend integration 已確認在 `frontend-AWS/src/api/visitors.js`，呼叫 `https://9u8ml80foj.execute-api.ap-northeast-1.amazonaws.com/views` 並讀取 response 的 `views` 欄位。",
                ],
                "bullets": [
                    "Lambda function 原始碼：To Confirm，目前 repo 中未找到 visitor counter Lambda implementation。",
                    "DynamoDB table name 與 key schema：To Confirm，目前 repo 文件只確認使用 DynamoDB 與 `/views` response。",
                    "IAM policy 細節：To Confirm，目前 repo 中未找到 AWS IAM policy 或 IaC 定義。",
                ],
            },
            {
                "heading": "4.2 GCP AI Backend",
                "page_break_before_table": True,
                "body": [
                    "Cloud Run backend 位於 `backend-GCP/`，Dockerfile 使用 `python:3.11-slim`，安裝 `requirements.txt` 後以 `uvicorn main:app --host 0.0.0.0 --port ${PORT}` 啟動。GitHub Actions 會 build Docker image、push 到 Artifact Registry，並透過 `gcloud run deploy` 部署。",
                    "FastAPI routes 分為 health、chat 與 rag 三類。`/ask-rag-stream` 使用 `StreamingResponse` 回傳 `text/event-stream`；`/ingest-docs` 受 `X-Admin-Token` 保護。",
                ],
                "table": {
                    "headers": ["Endpoint", "用途", "確認來源"],
                    "rows": [
                        ["GET `/`", "Health/config summary/startup warnings。", "`app/routes/health.py`"],
                        ["GET `/healthz`", "輕量 uptime health check。", "`app/routes/health.py`"],
                        ["POST `/chat`", "Basic Gemini chat，不使用文件 retrieval。", "`app/routes/chat.py`"],
                        ["POST `/chat-with-docs`", "讀取指定 GCS markdown 作為 direct context。", "`app/routes/chat.py`"],
                        ["POST `/ingest-docs`", "Admin-only 文件切 chunk、embedding、寫入 Firestore。", "`app/routes/rag.py`, `ingestion_service.py`"],
                        ["POST `/ask-rag`", "同步 RAG answer，回傳 answer、sources、session_id。", "`app/routes/rag.py`, `rag_service.py`"],
                        ["POST `/ask-rag-stream`", "Streaming RAG answer，回傳 metadata/token/done/error SSE。", "`app/routes/rag.py`, `rag_service.py`"],
                    ],
                },
            },
            {
                "heading": "4.3 RAG AI Assistant",
                "numbered": [
                    "文件來源：Cloud Storage bucket `cloud-resume-ai-rag-docs` 保存 markdown source documents；production notes 指出曾以 `CAPSTONE_PROJECT_STATE.md` 重建 index。",
                    "Chunking：`vector_service.chunk_text` 先依 Markdown headings 分段，再依 paragraph boundary 切分，最後才使用 size fallback。",
                    "Embedding：`gemini_service.embed_text` 使用 `text-embedding-005` 產生 query 與 chunk embeddings。",
                    "Retrieval：backend 目前 stream Firestore `document_chunks`，計算 cosine similarity，可選擇 hybrid keyword scoring、candidate pool、score threshold 與 deterministic reranking。",
                    "Response Generation：`rag_service` 將 top chunks 格式化為 `[S1]`, `[S2]` source context，要求 Gemini 的 factual claims 使用 source ID citations。",
                    "Persistence：backend 以 `session_id` 保存 user/assistant messages；query rewrite 若實際使用，會另存 `role: system`, `event_type: query_rewrite` 的 audit message。",
                ],
            },
        ],
    },
    {
        "heading": "5. 部署與維運",
        "subsections": [
            {
                "heading": "5.1 Frontend Deployment",
                "numbered": [
                    "GitHub Actions push to `main` 觸發 `Deploy Frontend to AWS` workflow。",
                    "Node.js 20 setup，使用 `npm ci` 安裝 dependencies。",
                    "寫入 `VITE_GCP_RAG_API_URL` 到 `.env`；若 secret 未提供，fallback 到目前 Cloud Run backend URL。",
                    "執行 `npm run build`。",
                    "使用 AWS credentials 透過 `aws s3 sync dist/ s3://... --delete` 部署。",
                    "透過 `aws cloudfront create-invalidation --paths \"/*\"` 清除 CDN cache。",
                ],
            },
            {
                "heading": "5.2 Backend Deployment",
                "numbered": [
                    "GitHub Actions push to `main` 且 path 符合 `backend-GCP/**` 或 backend workflow 時觸發。",
                    "透過 Google Workload Identity Federation 驗證到 GCP。",
                    "Build Docker image 並 push 到 Artifact Registry。",
                    "用 `gcloud run deploy` 部署 Cloud Run，設定 CORS、INGESTION_ADMIN_TOKEN 與 query rewrite env vars。",
                    "目前 workflow 中 `RAG_QUERY_REWRITE_ENABLED` 設為 `false`，代表 query rewriting 功能已實作但 deployment 預設未啟用。",
                ],
            },
            {
                "heading": "5.3 Logging",
                "body": [
                    "Backend 已加入 JSON-formatted stdout logs、request IDs、request duration logs 與 service boundary metadata logs。Cloud Run 可透過 Cloud Logging 查詢這些 stdout/stderr logs。",
                    "Lambda logs 應位於 CloudWatch Logs，但 repo 中沒有 Lambda 原始碼或 CloudWatch log group 設定；具體 log group name 為 To Confirm。",
                ],
            },
            {
                "heading": "5.4 Monitoring",
                "body": [
                    "目前已有 health endpoints、startup warnings、structured logging 與部分 request duration header。完整 monitoring dashboard、analytics、alerting 尚未完成，文件與 roadmap 均將其列為 future improvement。"
                ],
            },
        ],
    },
    {
        "heading": "6. 問題排除與技術決策",
        "subsections": [
            {
                "heading": "6.1 問題排除紀錄",
                "table": {
                    "headers": ["問題", "Root Cause", "解決方式", "Lesson Learned"],
                    "rows": [
                        ["Production CloudFront AI assistant failed", "Cloud Run CORS allowed localhost origins but missed the production CloudFront origin.", "Added `https://dvzu3s2gq6iw.cloudfront.net` to backend CORS defaults and deployment env vars; added regression test.", "Production browser failures can look like generic fetch errors; verify CORS preflight directly."],
                        ["Backend deploy failed after adding CORS env var", "`gcloud --set-env-vars` parsed comma-separated CORS origins as multiple key/value pairs.", "Used custom delimiter syntax: `--set-env-vars \"^|^CORS_ALLOWED_ORIGINS=...\"`.", "Comma-bearing env vars need explicit delimiter handling in gcloud deploy commands."],
                        ["Firestore conversations did not appear after feature work", "Cloud Run was still serving an older revision.", "Redeployed backend and verified active revision plus Firestore writes.", "Check serving revision before debugging application logic."],
                        ["RAG content was stale/incomplete in V1", "Indexed source content did not reflect current multi-cloud architecture.", "Updated GCS source to `CAPSTONE_PROJECT_STATE.md`, cleared stale chunks, and rebuilt index through `/ingest-docs`.", "RAG correctness depends on source freshness, not only prompt quality."],
                        ["Streaming fallback behavior", "Streaming can fail due to network/CORS/backend issues.", "Frontend tries `/ask-rag-stream` first and preserves `/ask-rag` fallback.", "Fallback paths help keep UX resilient while enabling streaming."],
                        ["DynamoDB reserved word issue", "To Confirm: not found in current repo Markdown or source files.", "To Confirm.", "Do not include this as a confirmed incident until the exact evidence is added."],
                        ["Vertex AI permission issue", "To Confirm: not found as a documented incident in current repo files.", "To Confirm.", "Permission troubleshooting should be documented with exact error text when it occurs."],
                    ],
                },
            },
            {
                "heading": "6.2 技術決策分析",
                "table": {
                    "headers": ["技術決策", "替代方案", "選擇原因", "Trade-off"],
                    "rows": [
                        ["AWS S3 + CloudFront", "EC2/Nginx, Amplify, Vercel", "適合 static portfolio，成本低，CDN/HTTPS 直觀，符合 Cloud Resume Challenge。", "動態行為需另外接 API；cache invalidation 需要部署流程處理。"],
                        ["API Gateway + Lambda + DynamoDB", "Container backend, EC2, RDS", "serverless visitor counter 簡潔、可展示 AWS event-driven pattern。", "Lambda/DynamoDB implementation details 目前未在 repo 中保存，後續可補 IaC 與 source。"],
                        ["Cloud Run", "Cloud Functions, GKE, Lambda", "FastAPI container 部署彈性高，適合 Python RAG backend 與 streaming endpoint。", "需要管理 container build、env vars、service account 與 CORS。"],
                        ["Vertex AI / Gemini", "OpenAI API, Bedrock, self-hosted model", "GCP pivot 讓 RAG backend 可較快完成 end-to-end deployed path。", "原 AWS Bedrock RAG 路線被 deferred，multi-cloud 故事需要清楚說明取捨。"],
                        ["Firestore", "Cloud SQL, BigQuery, vector DB", "簡化 document chunks 與 session conversation persistence。", "目前 retrieval 仍 full scan in memory，長期應升級 managed vector search。"],
                        ["Multi-Cloud", "AWS-only, GCP-only", "AWS 展示 serverless fundamentals，GCP 展示 AI/RAG engineering，形成更完整作品集。", "部署、CORS、文件與監控複雜度上升。"],
                        ["Terraform / full CI/CD 尚未完全導入", "手動部署或只用 GitHub Actions", "目前 repo 已有 GitHub Actions deployment，但基礎設施 IaC 尚未完整落地。", "短期可快速迭代；長期 reproducibility 與 auditability 仍需 Terraform/IaC 補強。"],
                    ],
                },
            },
        ],
    },
    {
        "heading": "7. 未來規劃",
        "subsections": [
            {
                "heading": "7.1 Terraform Infrastructure as Code",
                "body": ["將 AWS S3、CloudFront、API Gateway、Lambda、DynamoDB，以及 GCP Cloud Run、Artifact Registry、Firestore、GCS、IAM/service account 模組化管理。"],
            },
            {
                "heading": "7.2 GitHub Actions CI/CD",
                "body": ["Frontend/backend deployment workflows 已存在。下一步可加入 pull request checks、RAG evaluation gate、deployment smoke tests、Cloud Run revision verification 與 rollback notes。"],
            },
            {
                "heading": "7.3 Vector Search",
                "body": ["目前 Firestore chunks 是 in-memory full scan retrieval。未來可評估 Firestore Vector Search 或 Vertex AI Vector Search，將 retrieval 升級為 production-style ANN vector index。"],
            },
            {
                "heading": "7.4 Monitoring 與 Alerting",
                "body": ["使用 Cloud Logging / Cloud Monitoring 建立 request latency、error rate、RAG source usage、streaming error、Firestore read volume、Cloud Run revision health 等 dashboard 與 alerts。"],
            },
            {
                "heading": "7.5 Production Hardening",
                "bullets": [
                    "Authentication：保護管理端與 ingestion path，目前 `/ingest-docs` 已有 admin token。",
                    "Rate limiting：避免 public assistant 被濫用或造成成本暴增。",
                    "Better error handling：保留 controlled JSON/SSE errors，並增加 user-friendly classification。",
                    "Security hardening：補充 IAM least privilege、secret rotation、CORS origin review。",
                    "Cost monitoring：追蹤 Cloud Run、Vertex AI、Firestore reads、CloudFront 與 API Gateway 成本。",
                ],
            },
        ],
    },
    {
        "heading": "8. 技能總結",
        "table": {
            "headers": ["能力分類", "對應技術", "專案中如何體現"],
            "rows": [
                ["Cloud Engineering", "AWS, GCP, Cloud Run, S3, CloudFront", "將 portfolio delivery、serverless counter 與 AI backend 拆成可部署的 cloud paths。"],
                ["Serverless Architecture", "API Gateway, Lambda, DynamoDB, Cloud Run", "visitor counter 與 containerized FastAPI backend 都使用 managed/serverless execution model。"],
                ["AI Engineering", "Gemini, text-embedding-005, RAG", "實作 chunking、embedding、retrieval、source citations、streaming generation。"],
                ["Backend Development", "FastAPI, Python services/routes/schemas", "將 backend 拆成 config、schemas、routes、services，保留 `main:app` entrypoint。"],
                ["Frontend Development", "React, Vite, hooks, components", "模組化 portfolio UI、assistant streaming parsing、modal layout 與 bilingual content。"],
                ["DevOps", "GitHub Actions, Docker, gcloud, AWS CLI", "自動 build/deploy frontend 與 backend，處理 env vars 與 CloudFront invalidation。"],
                ["Security", "CORS, admin token, secrets", "修正 production CORS，保護 `/ingest-docs`，使用 GitHub secrets 傳入 deployment env vars。"],
                ["Troubleshooting", "CORS preflight, Cloud Run revisions, regression tests", "以 exact evidence 找到 live AI assistant failure root cause 並補測試。"],
                ["Technical Documentation", "Statement_MD, test records, development logs", "持續記錄 project state、RAG evolution、frontend changes、production incidents。"],
            ],
        },
    },
    {
        "heading": "9. 附錄",
        "subsections": [
            {
                "heading": "9.1 API Endpoints",
                "table": {
                    "headers": ["Endpoint", "Method", "用途"],
                    "rows": [
                        ["AWS `/views`", "GET", "Visitor count endpoint; frontend reads `views`."],
                        ["GCP `/`", "GET", "Backend health/config summary."],
                        ["GCP `/healthz`", "GET", "Lightweight health check."],
                        ["GCP `/chat`", "POST", "Basic Gemini chat."],
                        ["GCP `/chat-with-docs`", "POST", "Direct GCS document context chat."],
                        ["GCP `/ingest-docs`", "POST", "Admin-only ingestion into Firestore."],
                        ["GCP `/ask-rag`", "POST", "Synchronous RAG answer."],
                        ["GCP `/ask-rag-stream`", "POST", "Streaming RAG answer via SSE."],
                    ],
                },
            },
            {
                "heading": "9.2 Firestore Schema",
                "code": """```text
document_chunks/{deterministic_sha256(file_name:chunk_index)}
  file_name: string
  chunk_index: number
  chunk_text: string
  embedding: number[]
  content_hash: string
  char_count: number
  heading: string | null
  ingestion_key: string
  updated_at: server timestamp

conversations/{session_id}
  updated_at: server timestamp
  last_request_id: string | null

conversations/{session_id}/messages/{message_id}
  role: "user" | "assistant" | "system"
  content: string
  created_at: server timestamp
  request_id: string | optional
  event_type: "query_rewrite" | optional
  original_question: string | optional
  rewritten_query: string | optional
  rewrite_used: boolean | optional
```""",
            },
            {
                "heading": "9.3 DynamoDB Schema",
                "body": [
                    "To Confirm：目前 repo 文件與 frontend code 確認 visitor counter 使用 DynamoDB，且 `/views` response 回傳 `{ \"views\": number }`，但未找到 table name、partition key、attribute names 或 Lambda update expression。建議補上 Lambda source、IaC 或手動設定截圖/紀錄。"
                ],
            },
            {
                "heading": "9.4 Environment Variables",
                "table": {
                    "headers": ["變數", "位置", "用途"],
                    "rows": [
                        ["`VITE_GCP_RAG_API_URL`", "frontend-AWS `.env` / GitHub secret", "Frontend RAG backend base URL；缺省 fallback 到 Cloud Run URL。"],
                        ["`CORS_ALLOWED_ORIGINS`", "Cloud Run env", "允許 localhost 與 production CloudFront origin。"],
                        ["`INGESTION_ADMIN_TOKEN`", "Cloud Run env / GitHub secret", "保護 `/ingest-docs`。"],
                        ["`RAG_QUERY_REWRITE_ENABLED`", "Cloud Run env", "控制 query rewrite；目前 workflow 設為 `false`。"],
                        ["`RAG_QUERY_REWRITE_HISTORY_LIMIT`", "Cloud Run env", "query rewrite 使用的 recent history 筆數。"],
                        ["`RAG_QUERY_REWRITE_MODEL`", "Cloud Run env", "query rewrite model，預設 `gemini-2.5-flash`。"],
                        ["`GOOGLE_CLOUD_PROJECT`", "Cloud Run env", "Firestore/GCP client project config。"],
                        ["`GOOGLE_CLOUD_LOCATION`", "Cloud Run env", "Vertex AI location；settings default 為 `us-central1`。"],
                        ["`DOCS_BUCKET`", "Cloud Run env", "GCS markdown source bucket，default `cloud-resume-ai-rag-docs`。"],
                        ["`INGEST_DOCUMENTS`", "Cloud Run env", "ingestion source list；settings default 仍為 `PROJECT_STATE.md,Frontend_Development_Log.md`。"],
                        ["`DIRECT_CONTEXT_DOCUMENTS`", "Cloud Run env", "`/chat-with-docs` source list；settings default 仍為 `PROJECT_STATE.md,Frontend_Development_Log.md`。"],
                        ["`RAG_TOP_K`, `RAG_CANDIDATE_POOL_SIZE`, `RAG_SCORE_THRESHOLD`", "Cloud Run env", "retrieval result count、candidate pool 與 threshold。"],
                    ],
                },
            },
            {
                "heading": "9.5 Repository Structure",
                "code": """```text
.
├── Statement_MD/
│   ├── CAPSTONE_PROJECT_STATE.md
│   ├── CAPSTONE_V1_TEST_RECORD.md
│   ├── GCP_RAG_DEVELOPMENT_LOG.md
│   ├── GCP_RAG_PROJECT_STATE.md
│   └── REACT_Frontend_Development_Log.md
├── frontend-AWS/
│   ├── src/api/
│   ├── src/components/
│   ├── src/content/
│   ├── src/hooks/
│   └── src/pages/
├── backend-GCP/
│   ├── app/config/
│   ├── app/routes/
│   ├── app/schemas/
│   ├── app/services/
│   ├── Dockerfile
│   └── main.py
└── .github/workflows/
    ├── deploy-frontend.yml
    └── deploy-backend-gcp.yml
```""",
            },
            {
                "heading": "9.6 To Confirm 清單與不一致處理",
                "bullets": [
                    "AWS Lambda visitor counter source、DynamoDB table schema、IAM policy details 未在目前 repo 找到。",
                    "DynamoDB reserved word incident 與 Vertex AI permission incident 未在目前 repo 文件中找到具體證據。",
                    "`CAPSTONE_PROJECT_STATE.md` 與 `REACT_Frontend_Development_Log.md` 仍有部分 `frontend-Vite` 路徑，但目前 active app 與 workflow 是 `frontend-AWS`。",
                    "`frontend-AWS/src/content/portfolioContent.js` 的 hero/早期敘述仍出現 AWS Bedrock/S3 Vectors/Lambda 協調 RAG 的舊方向；目前 source-of-truth 文件說明 AI/RAG backend 已 pivot 到 GCP Cloud Run + Gemini + Firestore + GCS。",
                    "`backend-GCP/app/config/settings.py` default `INGEST_DOCUMENTS` / `DIRECT_CONTEXT_DOCUMENTS` 仍是 `PROJECT_STATE.md,Frontend_Development_Log.md`，但 production docs 說明目前 ingestion source 曾調整為 `CAPSTONE_PROJECT_STATE.md`；backend deployment workflow 沒有顯式設定這兩個 env vars。",
                    "Query rewriting 已在 code 中實作，但 `.github/workflows/deploy-backend-gcp.yml` 目前設定 `RAG_QUERY_REWRITE_ENABLED: \"false\"`，文件應描述為 available/implemented but disabled by current deploy config。",
                ],
            },
        ],
    },
]


def iter_blocks():
    yield ("title", TITLE)
    yield ("paragraph", SUBTITLE)
    for section in SECTIONS:
        yield from emit_section(section, level=1)


def emit_section(section, level):
    yield (f"h{level}", section["heading"])
    for para in section.get("body", []):
        yield ("paragraph", para)
    if "bullets" in section:
        for item in section["bullets"]:
            yield ("bullet", item)
    if "numbered" in section:
        for item in section["numbered"]:
            yield ("number", item)
    if "table" in section:
        yield ("table", section["table"])
    if "code" in section:
        yield ("code", section["code"])
    for subsection in section.get("subsections", []):
        yield from emit_section(subsection, level=min(level + 1, 3))


def markdown():
    lines = [f"# {TITLE}", "", SUBTITLE, ""]
    for section in SECTIONS:
        lines.extend(markdown_section(section, 2))
    MD_OUT.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def markdown_section(section, level):
    lines = ["#" * level + " " + section["heading"], ""]
    for para in section.get("body", []):
        lines.extend([para, ""])
    for item in section.get("bullets", []):
        lines.append(f"- {item}")
    if section.get("bullets"):
        lines.append("")
    for idx, item in enumerate(section.get("numbered", []), start=1):
        lines.append(f"{idx}. {item}")
    if section.get("numbered"):
        lines.append("")
    if "table" in section:
        table = section["table"]
        lines.append("| " + " | ".join(table["headers"]) + " |")
        lines.append("| " + " | ".join(["---"] * len(table["headers"])) + " |")
        for row in table["rows"]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    if "code" in section:
        lines.extend([section["code"], ""])
    for subsection in section.get("subsections", []):
        lines.extend(markdown_section(subsection, min(level + 1, 4)))
    return lines


def set_font(run, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run)
    run.font.size = Pt(9.5)
    run.bold = bold


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "DADCE0")
        borders.append(tag)
    tbl_pr.append(borders)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.font.size = Pt(8)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.0


def next_numbering_id(numbering):
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_id(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)

    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def render_docx_section(doc, section, level):
    doc.add_heading(section["heading"], level=level)

    for para in section.get("body", []):
        p = doc.add_paragraph(para)
        for run in p.runs:
            set_font(run)

    for item in section.get("bullets", []):
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_font(run)

    if section.get("numbered"):
        num_id = create_decimal_numbering(doc)
        for item in section["numbered"]:
            p = doc.add_paragraph(item)
            apply_numbering(p, num_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                set_font(run)

    if "table" in section:
        if section.get("page_break_before_table"):
            doc.add_page_break()
        add_docx_table(doc, section["table"])

    if "code" in section:
        for line in section["code"].splitlines():
            p = doc.add_paragraph(line, style="CodeBlock")
            for run in p.runs:
                set_font(run, "Courier New")

    for subsection in section.get("subsections", []):
        render_docx_section(doc, subsection, min(level + 1, 3))


def add_docx_table(doc, payload):
    headers = payload["headers"]
    rows = payload["rows"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)
    doc.add_paragraph()


def docx():
    doc = Document()
    style_doc(doc)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(TITLE)
    set_font(title_run)
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.bold = False

    subtitle = doc.add_paragraph(SUBTITLE)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    set_font(subtitle.runs[0])

    for section in SECTIONS:
        render_docx_section(doc, section, 1)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Multi-Cloud AI Portfolio Assistant technical document"
    doc.core_properties.keywords = "AWS, GCP, RAG, Cloud Run, Gemini, Firestore, CloudFront"
    doc.save(DOCX_OUT)


if __name__ == "__main__":
    markdown()
    docx()
    print(MD_OUT)
    print(DOCX_OUT)
